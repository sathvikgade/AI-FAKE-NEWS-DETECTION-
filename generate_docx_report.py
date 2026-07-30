import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_ieee_docx():
    doc = docx.Document()

    # Set Margins (IEEE Standard 0.75 in / 1 in)
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styling helper
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(15, 23, 42)

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(100, 116, 139)

    def add_h1(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)

    def add_h2(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 65, 85)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    def add_p(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        return p

    # Document Header
    add_title("AI-Powered Fake News Detection Using Text Classification")
    add_subtitle("IEEE Standard Documentation Report | Summer Internship Program in AI & ML (2026)")
    add_subtitle("Indian Institute of Computing and Technology (IICT)")
    doc.add_paragraph()

    # Abstract Box
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.left_indent = Inches(0.3)
    p_abs.paragraph_format.right_indent = Inches(0.3)
    r_abs_t = p_abs.add_run("Abstract— ")
    r_abs_t.bold = True
    r_abs_t.font.name = 'Arial'
    r_abs_t.font.size = Pt(10)
    r_abs = p_abs.add_run("The rapid proliferation of digital media and social networks has accelerated the spread of misinformation, making automated fake news detection a crucial task for natural language processing (NLP) and machine learning (ML). This paper presents a complete machine learning pipeline implemented from scratch to classify news articles as Real or Fake. We perform systematic text preprocessing, manual tokenization, stop-words removal, and Term Frequency-Inverse Document Frequency (TF-IDF) feature extraction. Four distinct classification paradigms—K-Nearest Neighbors (Non-Parametric), Logistic Regression (Parametric), Random Forest (Ensemble), and Multi-Layer Perceptron (Neural Network)—are evaluated on a dataset of 39,999 news articles. Experimental results demonstrate that ensemble learning via Random Forest achieves a top accuracy of 99.72% and F1-score of 100.00%, outperforming traditional non-parametric distance-based models which suffer from the curse of dimensionality.")
    r_abs.font.name = 'Calibri'
    r_abs.font.size = Pt(10)
    r_abs.italic = True

    add_p("Keywords— Fake News Detection, Machine Learning, Text Classification, TF-IDF, Random Forest, Logistic Regression, Neural Networks, IEEE Standard.")

    # 1. Introduction
    add_h1("1. Introduction")
    add_h2("1.1 Problem Statement")
    add_p("In contemporary digital ecosystems, misinformation and fake news propagate rapidly across online news portals and social platforms. Fake news refers to intentionally fabricated stories presented as authentic journalism to deceive readers, manipulate public opinion, or generate ad revenue. Manual verification by human fact-checkers is unscalable due to the high volume of daily publication. Thus, building an automated machine learning system capable of analyzing text structure, vocabulary, and stylistic markers to distinguish real news from fake news is essential.")

    add_h2("1.2 Importance of Fake News Detection")
    add_p("Unfiltered fake news poses significant threats to democratic processes, financial markets, public health safety, and social harmony. Automated NLP-driven text classification offers real-time verification capabilities, empowering platforms to flag suspicious content before viral dissemination occurs.")

    add_h2("1.3 30-Day Workflow Objectives")
    add_p("Following the IICT Summer Internship curriculum:\n"
          "• Week 1: Data acquisition, text cleaning (punctuation and stop-words removal), manual tokenization.\n"
          "• Week 2: Feature engineering via Bag-of-Words / TF-IDF vectorization and exploratory data analysis.\n"
          "• Week 3: Algorithm implementation across KNN, Logistic Regression, Random Forest, and Simple Neural Network (MLP).\n"
          "• Week 4: Quantitative model evaluation (Accuracy, Precision, Recall, F1-Score), confusion matrix visualization, and IEEE report documentation.")

    # 2. Dataset Description
    add_h1("2. Dataset Description")
    add_h2("2.1 Dataset Source and Distribution")
    add_p("The project utilizes the Kaggle Fake and Real News Dataset, comprising 39,999 news articles collected from verified news agencies and flagged misinformation outlets.")

    # Table 1
    table1 = doc.add_table(rows=4, cols=4)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers1 = ["Class Label", "Category", "Sample Count", "Percentage"]
    for i, h in enumerate(headers1):
        cell = table1.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    data1 = [
        ["0", "Fake News", "20,886", "52.2%"],
        ["1", "Real News", "19,113", "47.8%"],
        ["Total", "Full Dataset", "39,999", "100.0%"]
    ]
    for row_idx, row_data in enumerate(data1, start=1):
        for col_idx, text in enumerate(row_data):
            table1.cell(row_idx, col_idx).text = text

    add_h2("2.2 Dataset Attributes")
    add_p("Each dataset record contains six key attributes:\n"
          "1. index: Unique identifier.\n"
          "2. title: Headline of the news article.\n"
          "3. text: Full body text of the article.\n"
          "4. subject: Content topic (e.g., politics, world news).\n"
          "5. date: Publication timestamp.\n"
          "6. class / label: Binary classification target (Fake or Real).")

    # 3. Methodology
    add_h1("3. Methodology")
    add_h2("3.1 Text Preprocessing Pipeline")
    add_p("Raw textual data contains noise, punctuation, capitalizations, and non-informative frequent words (stop-words). We implement a multi-stage cleaning function:\n"
          "1. Agency Marker Normalization: Strip static agency prefixes such as 'WASHINGTON (Reuters) -' to prevent shortcut learning.\n"
          "2. Lowercasing: Convert all characters to lowercase.\n"
          "3. Regex Punctuation Removal: Remove punctuation, special characters, and numbers using regex pattern [^a-z\\s].\n"
          "4. Manual Tokenization: Split text into individual word tokens based on whitespace.\n"
          "5. Stop-words Removal: Filter out non-discriminatory English stop-words (the, is, at, which) using NLTK corpus dictionaries.")

    add_h2("3.2 Feature Extraction (TF-IDF)")
    add_p("We map cleaned text strings into a continuous numerical vector space using Term Frequency-Inverse Document Frequency (TF-IDF) vectorization, extracting unigrams and bigrams bound to the top 5,000 to 10,000 feature dimensions.")

    add_h2("3.3 Classification Algorithms")
    add_p("• K-Nearest Neighbors (KNN - Non-Parametric): Distance-based classification using 5 nearest neighbors.\n"
          "• Logistic Regression (Parametric Linear): Sigmoid activation function over linear decision boundaries.\n"
          "• Random Forest (Ensemble): Combination of 100 de-correlated decision trees with bootstrap aggregation.\n"
          "• Simple Neural Network (MLP): Multi-Layer Perceptron with 100 hidden neurons and ReLU activation.")

    # 4. Results
    add_h1("4. Experimental Results")
    add_p("We evaluate all models across 8,000 holdout test samples using standard evaluation metrics:")

    # Table 2
    table2 = doc.add_table(rows=5, cols=6)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ["Algorithm", "Model Paradigm", "Accuracy", "Precision", "Recall", "F1-Score"]
    for i, h in enumerate(headers2):
        cell = table2.cell(0, i)
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)

    data2 = [
        ["KNN", "Non-Parametric", "76.46%", "91.00%", "57.00%", "70.00%"],
        ["Logistic Regression", "Parametric Linear", "98.67%", "98.00%", "99.00%", "99.00%"],
        ["Neural Network (MLP)", "Deep Learning", "99.09%", "99.00%", "99.00%", "99.00%"],
        ["Random Forest", "Ensemble Trees", "99.72%", "100.00%", "100.00%", "100.00%"]
    ]
    for row_idx, row_data in enumerate(data2, start=1):
        for col_idx, text in enumerate(row_data):
            cell = table2.cell(row_idx, col_idx)
            cell.text = text
            if row_idx == 4: # Highlight Random Forest
                cell.paragraphs[0].runs[0].font.bold = True

    # 5. Discussion
    add_h1("5. Discussion")
    add_h2("5.1 Parametric vs. Non-Parametric Model Comparison")
    add_p("• Parametric Models (Logistic Regression & Neural Net): Summarize data through fixed weights. Logistic Regression achieved 98.67% accuracy with extremely fast training (< 2 seconds), proving linear decision boundaries are highly effective for high-dimensional sparse text vectors.\n"
          "• Non-Parametric Models (KNN): Achieved lower accuracy (76.46%) and recall (57.00%) due to the Curse of Dimensionality in 5,000-dimensional TF-IDF space.\n"
          "• Ensemble Superiority (Random Forest): Achieved the top accuracy (99.72%) by combining 100 decision trees, cancelling out individual variance.")

    # 6. Conclusion
    add_h1("6. Conclusion & Future Scope")
    add_p("• Key Insights: Text cleaning and TF-IDF feature extraction enable near-perfect text classification performance.\n"
          "• Limitations: Cross-domain performance decay on new topics outside the training period.\n"
          "• Future Scope: Integration of Transformer models (BERT, RoBERTa, LLMs) and multimodal graph detection.")

    # 7. Appendix
    add_h1("7. Appendix: Test Samples & Code")
    add_p("The full Python code implementations (fake_news_pipeline.py, preprocessing.py, feature_engineering.py, train.py, evaluate.py, predict.py) and dataset test suites are hosted on GitHub:\n"
          "https://github.com/sathvikgade/AI-FAKE-NEWS-DETECTION-.git")

    # 8. References
    add_h1("8. References (IEEE Format)")
    add_p("[1] H. Ahmed, I. Traore, and S. Saad, 'Detecting fake news headlines with machine learning,' IEEE Access, vol. 6, pp. 62141–62153, 2018.\n"
          "[2] V. Perez-Rosas et al., 'Automatic detection of fake news,' in Proc. COLING, 2018, pp. 3391–3401.\n"
          "[3] L. Breiman, 'Random forests,' Machine Learning, vol. 45, no. 1, pp. 5–32, 2001.")

    output_path = r"C:\Users\LENOVO\OneDrive\Desktop\AI_FAKE_NEWS_PROJECT\IEEE_Fake_News_Detection_Report.docx"
    doc.save(output_path)
    print(f"IEEE Word Document generated successfully at: {output_path}")

if __name__ == "__main__":
    create_ieee_docx()
